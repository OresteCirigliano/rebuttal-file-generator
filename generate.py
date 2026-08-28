import os
import sys
import glob
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx import Document

# 1. Legge i parametri passati da GitHub Actions
customer_name = sys.argv[1]
subscription_id = sys.argv[2]
subscription_creation_date = sys.argv[3]
shipping_order_id = sys.argv[4]
tracking_id = sys.argv[5]
shipping_total_value = sys.argv[6]
disputed_amount = sys.argv[7]
delivery_date = sys.argv[8]
consecutive_shipment_number = sys.argv[9]
disputed_charge_date = sys.argv[10]
disputed_instalment = sys.argv[11]
chargeback_reason = sys.argv[12]

# 2. Calcolo automatico delle date degli instalment (+1, +2, +3 mesi)
instalment_1_date = subscription_creation_date
instalment_2_date = "N/A"
instalment_3_date = "N/A"
instalment_4_date = "N/A"

date_formats = (
    "%d.%m.%Y", "%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y",
    "%d.%m.%Y %H:%M", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"
)

clean_date_str = subscription_creation_date.split()[0] if subscription_creation_date else ""

for date_format in date_formats:
    try:
        base_date = datetime.strptime(clean_date_str, date_format.split()[0])
        instalment_1_date = base_date.strftime("%d.%m.%Y")
        instalment_2_date = (base_date + relativedelta(months=1)).strftime("%d.%m.%Y")
        instalment_3_date = (base_date + relativedelta(months=2)).strftime("%d.%m.%Y")
        instalment_4_date = (base_date + relativedelta(months=3)).strftime("%d.%m.%Y")
        break
    except ValueError:
        continue

# 3. Ricerca del file template .docx
docx_files = glob.glob("*.docx")
if not docx_files:
    raise FileNotFoundError("Nessun file .docx trovato!")

template_path = docx_files[0]
doc = Document(template_path)

# 4. Mappatura completa di tutti i segnaposto
replacements = {
    "{CUSTOMER FULL NAME}": customer_name,
    "{SUBSCRIPTION ID}": subscription_id,
    "{SUBSCRIPTION DATE}": subscription_creation_date,
    "{ORDER ID}": shipping_order_id,
    "{TRACKING ID}": tracking_id,
    "{Shipping Value}": shipping_total_value,
    "{DISPUTE AMOUNT + CURRENCY}": disputed_amount,
    "{DELIVERY DATE}": delivery_date,
    "{shipping number}": consecutive_shipment_number,
    "{Disputed Charge Date}": disputed_charge_date,
    "{#X of 4}": disputed_instalment,
    "{e.g. Fraudulent / Not Received / Unauthorised}": chargeback_reason,
    "{INSTALMENT_1_DATE}": instalment_1_date,
    "{INSTALMENT_2_DATE}": instalment_2_date,
    "{INSTALMENT_3_DATE}": instalment_3_date,
    "{INSTALMENT_4_DATE}": instalment_4_date,
}

# 5. Sostituzione nei paragrafi Word
for p in doc.paragraphs:
    for key, value in replacements.items():
        if key in p.text:
            p.text = p.text.replace(key, value)

# 6. Sostituzione nelle tabelle Word
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for key, value in replacements.items():
                    if key in p.text:
                        p.text = p.text.replace(key, value)

# 7. Salva il documento Word aggiornato
output_filename = f"Rebuttal_{subscription_id}_{customer_name.replace(' ', '_')}.docx"
doc.save(output_filename)
print(f"File Word generato con successo: {output_filename}")

# 8. Generazione del testo esplicativo per la banca (.txt e console)
context_text = f"""--- BANK DISPUTE SUMMARY & CONTEXT ---

On {subscription_creation_date}, the customer ({customer_name}) placed an ongoing instalment order (Subscription ID: {subscription_id}). 
The cardholder explicitly opted to receive a 4-month supply shipment (Shipment #{consecutive_shipment_number}, Order ID: {shipping_order_id}) and selected the instalment plan option to divide the total shipping value ({shipping_total_value}) into 4 payments ({disputed_amount} per instalment).

During checkout, the customer formally accepted the Terms & Conditions (mandatory to proceed with the payment) and authenticated the initial charge via 3D-Secure (3DS). 
The full product supply was successfully delivered on {delivery_date} via DHL (Tracking ID: {tracking_id}). 

The customer is currently disputing instalment {disputed_instalment} for reason '{chargeback_reason}'. 
This charge is fully legitimate and contractually due, as it partially covers inventory delivered on {delivery_date} and retained by the customer.
"""

print("\n" + context_text)

# Salva il testo anche in un file TXT pronto per il download
txt_filename = f"Context_{subscription_id}.txt"
with open(txt_filename, "w", encoding="utf-8") as f:
    f.write(context_text)

print(f"File di testo salvato con successo: {txt_filename}")
