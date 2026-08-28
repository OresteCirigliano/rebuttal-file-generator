import streamlit as st
import docx
from datetime import datetime
from dateutil.relativedelta import relativedelta
import io
import zipfile

st.set_page_config(page_title="Chargeback Rebuttal Generator", page_icon="📄")

st.title("📄 Chargeback Rebuttal Generator")
st.markdown("Compila i campi sottostanti per generare la documentazione difensiva completa (.docx + .txt).")

# Form di inserimento dati
with st.form("rebuttal_form"):
    col1, col2 = st.columns(2)
    
    with col1:
        customer_name = st.text_input("Customer Full Name", value="John Doe")
        subscription_id = st.text_input("Subscription ID", value="SUB-12345")
        subscription_creation_date = st.text_input("Subscription Creation Date (DD.MM.YYYY)", value="26.06.2026")
        shipping_order_id = st.text_input("Shipping Order ID", value="ORD-98765")
        tracking_id = st.text_input("Tracking ID", value="DHL123456789")
        shipping_total_value = st.text_input("Shipping Total Value", value="£120.00 EUR")

    with col2:
        disputed_amount = st.text_input("Disputed Amount + Currency", value="£30.00 EUR")
        delivery_date = st.text_input("Delivery Date (DD.MM.YYYY)", value="28.06.2026")
        consecutive_shipment_number = st.text_input("Consecutive Shipment Number", value="1")
        disputed_charge_date = st.text_input("Disputed Charge Date (DD.MM.YYYY)", value="26.07.2026")
        disputed_instalment = st.text_input("Disputed Instalment (#X of 4)", value="2")
        chargeback_reason = st.selectbox("Chargeback Reason", ["Unauthorised / Fraudulent", "Goods Not Received", "Subscription Cancelled", "Product unacceptable", "Credit not processed", "Other"])

    submitted = st.form_submit_button("🚀 Genera Rebuttal Package")

if submitted:
    # 1. Calcolo automatico delle date degli instalment (+1, +2, +3 mesi)
    instalment_1_date = subscription_creation_date
    instalment_2_date = "N/A"
    instalment_3_date = "N/A"
    instalment_4_date = "N/A"

    date_formats = ("%d.%m.%Y", "%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y")
    clean_date_str = subscription_creation_date.split()[0] if subscription_creation_date else ""

    for df in date_formats:
        try:
            base_date = datetime.strptime(clean_date_str, df)
            instalment_1_date = base_date.strftime("%d.%m.%Y")
            instalment_2_date = (base_date + relativedelta(months=1)).strftime("%d.%m.%Y")
            instalment_3_date = (base_date + relativedelta(months=2)).strftime("%d.%m.%Y")
            instalment_4_date = (base_date + relativedelta(months=3)).strftime("%d.%m.%Y")
            break
        except ValueError:
            continue

    # 2. Compilazione documento Word
    doc = docx.Document("template.docx")
    
    replacements = {
        "{CUSTOMER FULL NAME}": customer_name,
        "{SUBSCRIPTION ID}": subscription_id,
        "{SUBSCRIPTION DATE}": subscription_creation_date,
        "{ORDER ID}": shipping_order_id,
        "{TRACKING ID}": tracking_id,
        "{Shipping Value}": shipping_total_value,
        "{DISPUTE AMOUNT + CURRENCY}": disputed_amount,
        "{AMOUNT + CURRENCY}": disputed_amount,
        "{DELIVERY DATE}": delivery_date,
        "{shipping number}": consecutive_shipment_number,
        "{Disputed Charge Date}": disputed_charge_date,
        "{#X of 4}": disputed_instalment,
        "{e.g. Fraudulent / Not Received / Unauthorised}": chargeback_reason,
        "{INSTALMENT_1_DATE}": instalment_1_date,
        "{INSTALMENT_2_DATE}": instalment_2_date,
        "{INSTALMENT_3_DATE}": instalment_3_date,
        "{INSTALMENT_4_DATE}": instalment_4_date,
    }

    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in p.text:
                            p.text = p.text.replace(key, value)

    # Salva il Word in memoria
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    # 3. Generazione testo di contesto per la banca
    context_text = f"""--- BANK DISPUTE SUMMARY & CONTEXT ---

On {subscription_creation_date}, the customer ({customer_name}) placed an ongoing instalment order (Subscription ID: {subscription_id}). 
The cardholder explicitly opted to receive a 4-month supply shipment (Shipment #{consecutive_shipment_number}, Order ID: {shipping_order_id}) and selected the instalment plan option to divide the total shipping value ({shipping_total_value}) into 4 payments ({disputed_amount} per instalment).

During checkout, the customer formally accepted the Terms & Conditions (mandatory to proceed with the payment) and authenticated the initial charge via 3D-Secure (3DS). 
The full product supply was successfully delivered on {delivery_date} via DHL (Tracking ID: {tracking_id}). 

The customer is currently disputing instalment {disputed_instalment} for reason '{chargeback_reason}'. 
This charge is fully legitimate and contractually due, as it partially covers inventory delivered on {delivery_date} and retained by the customer.
"""

    # 4. Creazione pacchetto ZIP da far scaricare
    zip_io = io.BytesIO()
    with zipfile.ZipFile(zip_io, 'w') as zf:
        zf.writestr(f"Rebuttal_{subscription_id}.docx", docx_io.getvalue())
        zf.writestr(f"Context_{subscription_id}.txt", context_text)
        screenshots = [
            "Screenshot_T&C_5_ES.jpg",
            "Screenshot_T&C_5_UK.jpg",
            "Screenshot_T&C_8_ES.jpg",
            "Screenshot_T&C_8_UK.jpg",
            "Screenshot_checkout ESjpg.jpg",
            "Screenshot_checkout UK.jpg"
        ]
        
        # Ciclo per allegare ogni screenshot se presente nella cartella
        for img in screenshots:
            try:
                with open(img, "rb") as f:
                    zf.writestr(img, f.read())
            except FileNotFoundError:
                st.warning(f"⚠️ Immagine non trovata nel repository: {img}")
    zip_io.seek(0)

    st.success("✅ Documentazione generata con successo!")
    
    # Bottone di download dello ZIP
    st.download_button(
        label="📥 Scarica Pacchetto (.ZIP)",
        data=zip_io,
        file_name=f"Rebuttal_Package_{subscription_id}.zip",
        mime="application/zip"
    )
    
    # Anteprima testo di contesto
    st.subheader("Anteprima Testo per la Banca")
    st.code(context_text, language="text")
